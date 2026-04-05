import re
import streamlit as st
import docx
import google.generativeai as genai


# =========================
# PAGE CONFIG
# =========================
st.set_page_config(
    page_title="Word to MDD Converter",
    page_icon="📝",
    layout="wide"
)

st.title("📝 Word to MDD/Dimensions Script Converter")
st.markdown("Încarcă un fișier Word cu chestionarul tău.")


# =========================
# API KEY INPUT
# =========================
api_key = st.text_input("Introdu cheia ", type="password")


# =========================
# HELPERS
# =========================

# Instruction paragraph detection
INSTRUCTION_STARTERS = (
    "Bitte ", "Mehrfachantworten", "Mehrfach ",
)


def is_instruction_paragraph(para) -> bool:
    """
    Returns True if this paragraph is a survey instruction:
    - starts with 'Bitte ' or 'Mehrfachantworten'
    - AND is NOT fully bold (all-bold = question text, not instruction)
    """
    text = para.text.strip()
    if not text:
        return False
    if not any(text.startswith(s) for s in INSTRUCTION_STARTERS):
        return False
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    all_bold = all(r.bold is True for r in runs)
    return not all_bold  # instruction = NOT all-bold


def extract_runs_text(paragraph) -> str:
    """
    Extract text from a paragraph preserving bold (<b>) and underline (<u>).
    Only marks formatting when explicitly True — not None/False.

    Also handles the case where a single paragraph contains a \n followed by
    a "Bitte..." run that is not bold — in that case the Bitte part is marked
    [INSTRUCTION] so the LLM puts it in <span class=""instruction"">.
    """
    runs = paragraph.runs
    parts = []
    after_newline = False  # True if the previous meaningful run was a bare \n

    for run in runs:
        text = run.text
        if not text:
            continue

        # Case 1: run is purely a newline separator
        if text == "\n":
            after_newline = True
            parts.append("\n")
            continue

        # Case 2: run contains an embedded newline (e.g. "question\nBitte...")
        if "\n" in text:
            segments = text.split("\n")
            for s_i, seg in enumerate(segments):
                if s_i == 0:
                    t = seg
                    if run.bold is True and t:
                        t = f"<b>{t}</b>"
                    if run.underline is True and t:
                        t = f"<u>{t}</u>"
                    if t:
                        parts.append(t)
                else:
                    stripped = seg.strip()
                    if stripped and any(stripped.startswith(s) for s in INSTRUCTION_STARTERS) and run.bold is not True:
                        parts.append(f"\n[INSTRUCTION] {stripped}")
                    elif stripped:
                        t = stripped
                        if run.bold is True:
                            t = f"<b>{t}</b>"
                        if run.underline is True:
                            t = f"<u>{t}</u>"
                        parts.append(f"\n{t}")
            after_newline = False
            continue

        # Case 3: normal text run — check if it follows a bare \n and is an instruction
        stripped = text.strip()
        if after_newline and stripped and any(stripped.startswith(s) for s in INSTRUCTION_STARTERS) and run.bold is not True:
            # Replace the trailing \n in parts with [INSTRUCTION] marker
            if parts and parts[-1] == "\n":
                parts[-1] = ""  # remove the bare newline
            parts.append(f"\n[INSTRUCTION] {stripped}")
            after_newline = False
            continue

        after_newline = False
        if run.bold is True:
            text = f"<b>{text}</b>"
        if run.underline is True:
            text = f"<u>{text}</u>"
        parts.append(text)

    return "".join(parts)


def extract_cell_content(cell) -> str:
    """
    Extract all paragraphs from a cell.
    If the LAST paragraph looks like an instruction (starts with 'Bitte...'
    and is not all-bold), mark it with [INSTRUCTION] prefix.
    This lets the LLM know to put it in <span class=""instruction"">.
    """
    paras = [p for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return ""
    parts = []
    for i, para in enumerate(paras):
        is_last = (i == len(paras) - 1)
        if is_last and len(paras) > 1 and is_instruction_paragraph(para):
            parts.append(f"[INSTRUCTION] {para.text.strip()}")
        else:
            parts.append(extract_runs_text(para).strip())
    return "\n".join(p for p in parts if p)


def is_base_row(row) -> bool:
    """Returns True if this table row is a Base/routing row (row 0 with Base: text)."""
    for cell in row.cells:
        if cell.text.strip().startswith("Base:"):
            return True
    return False


def extract_base_text(row) -> str:
    """Extract the Base: text from a routing row, deduplicated."""
    seen = set()
    parts = []
    for cell in row.cells:
        txt = cell.text.strip()
        if txt and txt not in seen:
            seen.add(txt)
            parts.append(txt)
    return " ".join(parts)


def extract_text_from_docx(file) -> str:
    doc = docx.Document(file)
    content = []

    # Paragraphs — preserve bold/underline formatting
    for para in doc.paragraphs:
        text = extract_runs_text(para).strip()
        if text:
            content.append(text)

    # Tables — detect Base rows, instructions, preserve bold/underline
    table_lines = []
    for table in doc.tables:
        for r_i, row in enumerate(table.rows):
            # Row 0 with Base: → emit as [BASE] comment for MDD routing
            if r_i == 0 and is_base_row(row):
                base_text = extract_base_text(row)
                table_lines.append(f"[BASE] {base_text}")
                continue

            row_data = []
            for cell in row.cells:
                cell_text = extract_cell_content(cell)
                if cell_text:
                    row_data.append(cell_text)

            # Deduplicate consecutive identical cells
            unique_data = []
            for item in row_data:
                if item not in unique_data:
                    unique_data.append(item)

            if unique_data:
                table_lines.append(" | ".join(unique_data))

    if table_lines:
        content.append("\n--- TABELE ---")
        content.extend(table_lines)

    return "\n".join(content).strip()


def clean_model_output(text: str) -> str:
    if not text:
        return ""

    cleaned = text.strip()
    cleaned = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    return cleaned.strip()


def get_response_text(response) -> str:
    try:
        if hasattr(response, "text") and response.text:
            return response.text
    except Exception:
        pass

    try:
        parts = []
        for candidate in response.candidates:
            for part in candidate.content.parts:
                if hasattr(part, "text") and part.text:
                    parts.append(part.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""


def build_prompt(document_text: str) -> str:
    return f"""
You are an Expert Survey Programmer writing MDD/Dimensions script.
Convert the raw survey text into COMPLETE MDD code.
Do NOT summarize.
Do NOT explain.
Output ONLY raw MDD code.

CRITICAL SYNTAX RULES (DO NOT BREAK):

1. IGNORE SECTION HEADERS:
   Ignore non-question headers and noise such as:
   - "LINEARES TV"
   - "STREAMING"
   - page numbers
   - tester notes
   - wave notes
   - "not relevant for scripting"
   - background notes
   Only process actual survey questions and their answer structures.

2. DO NOT ADD HIDDEN QUESTIONS:
   Do NOT create hidden questions such as StatusCP.
   Do NOT invent helper variables, counters, or metadata questions unless they are explicit survey questions in the source.

3. PRESERVE ORIGINAL CODES EXACTLY:
   NEVER renumber, normalize, compress, or reorder answer codes.
   Keep original numeric codes exactly as they appear in the source, including non-sequential codes like 14, 21, 24, 99.

4. QUOTES ESCAPING:
   NEVER use backslashes (\\" ) for HTML.
   You MUST use double-double quotes.
   CORRECT: <div class=""qtext"">
   WRONG: <div class=\\"qtext\\">

5. SELECTIONS:
   - Use categorical [0..1] for Single Choice.
   - Use categorical [0..] or categorical [0..3] for Multiple Choice.
   - NEVER append max(3) or similar text after the categorical block.

6. OTHER / OPEN TEXT MUST STAY ON THE SAME LINE:
   For "Sonstiges" / "Other" style options, use the exact syntax:
   other(other_id "" text[0..]) fix
   and keep it on the SAME LINE as the statement.
   CORRECT:
   _13 "Sonstiges, und zwar:" other(other_13 "" text[0..]) fix

7. EXCLUSIVE CODES & RANDOMIZATION:
   Add fix exclusive to options such as:
   - "weiß nicht"
   - "keines der genannten"
   - "keinen davon"
   NEVER use ran except.
   Just use ran; at the end of the block.

8. GRID SYNTAX:
   The Question Name and Text MUST come BEFORE the loop keyword.
   DO NOT output SkipSecurity = "Yes".

9. GRID FIELDS:
   - For normal categorical grids, ALWAYS use _scale "" inside the fields - block.
   - NEVER invent field names like _providers "".
   - For simple open-end list grids, also use _scale "" but with text[0..]; instead of categorical.

10. PIPING / VARIABLES:
   Replace placeholders like:
   - [show kids name]
   - [show kids name in blue letters]
   with exact MDD piping syntax:
   {{{{#kids_name.response.value}}}}
   If color is specified, use:
   <span style='color:blue'><strong>{{{{#kids_name.response.value}}}}</strong></span>

11. BASE / ROUTING COMMENTS:
   The source document marks routing information with [BASE] at the start of a line.
   Example: "[BASE] Base: if Q2 item 2 ≠ code 6 (nie); [M]; [SC]"

   RULE: Output EVERY [BASE] line as an MDD comment on the line IMMEDIATELY BEFORE
   the question it belongs to, using single-quote syntax:
   'Base: if Q2 item 2 ≠ code 6 (nie); [M]; [SC]

   - Strip the [BASE] prefix, keep everything after it verbatim.
   - Place the comment directly before the question definition (no blank line between them).
   - DO NOT output routing logic (VisibleIf, filter conditions, masks) — only the comment.

   CORRECT:
   'Base: if Q2 item 2 (Mediatheken) ≠ code 6 (nie); [S per row]; [SC]
   Q7 "<div class=""qtext"">...</div>"
   ...

   WRONG (comment missing):
   Q7 "<div class=""qtext"">...</div>"

12. BANKED GRIDS:
   ALL normal categorical grids/loops MUST include:
   [
       GfKGridType = "banked"
   ]
   EXCEPTION:
   Do NOT use banked for simple repeated OE/title-entry grids.

13. HTML FORMATTING FOR QUESTION TEXTS:
   - The source text uses <b>...</b> to mark bold runs and <u>...</u> to mark underline runs.
   - These tags appear ONLY where the Word document has explicit formatting.
   - MANDATORY: Convert EVERY <b>text</b> from source → <strong>text</strong> in MDD output.
   - MANDATORY: Convert EVERY <u>text</u> from source → <u>text</u> in MDD output.
   - Do NOT skip, drop, or ignore any <b> or <u> tags from the source.
   - Do NOT add bold or underline where it is NOT in the source.
   - Do NOT wrap the entire question text in <strong> unless every run is marked <b>.
   - Keep all question text inside: <div class=""qtext""> ... </div>

   BOLD CONVERSION EXAMPLES:
   Source: "<b>Was ist die </b><u><b>maximale</b></u><b> Anzahl...</b>"
   MDD:    "<div class=""qtext""><strong>Was ist die </strong><u><strong>maximale</strong></u><strong> Anzahl...</strong></div>"

   Source: "Im Folgenden sehen Sie..."  (no <b> tags)
   MDD:    "<div class=""qtext"">Im Folgenden sehen Sie...</div>"  (no bold added)

14. HTML FORMATTING FOR INSTRUCTIONS — ALL QUESTION TYPES:

   HOW TO DETECT INSTRUCTIONS IN THE SOURCE:
   The source document marks instructions with the tag [INSTRUCTION] at the start of a line.
   Example source line: "[INSTRUCTION] Bitte wählen Sie bis zu drei Anbieter aus."
   This means the preprocessor already identified that paragraph as an instruction.

   RULE A — [INSTRUCTION] TAG PRESENT:
   - If a line in the source starts with [INSTRUCTION], ALWAYS output it as:
     <span class=""instruction"">...text without [INSTRUCTION] prefix...</span>
   - This applies to ALL question types: single, multi, grid (banked), OE, numeric.
   - The span MUST be on the NEXT LINE after </div>, never on the same line.
   - No space between </div> and the newline.

   RULE B — NO [INSTRUCTION] TAG:
   - If no [INSTRUCTION] tag appears in the source for a question, output ONLY:
     <div class=""qtext"">...</div>"
   - DO NOT invent any instruction text.
   - DO NOT add <span class=""instruction""> if [INSTRUCTION] is absent.

   RULE C — INSTRUCTION INSIDE QUESTION TEXT:
   - Sometimes an instruction sentence appears at the END of the question text paragraph
     (e.g. "...am ehesten? Bitte wählen Sie bis zu drei Anbieter aus.").
   - In that case: split the text. Put the question part in <div class=""qtext"">,
     and move the "Bitte..." sentence to <span class=""instruction"">.
   - A sentence is an instruction if it starts with "Bitte " and comes AFTER the main question.

   FORMATTING — ALWAYS:
   - </div> immediately followed by newline, then <span class=""instruction""> on the next line.
   - NEVER: </div><span ...> (same line)
   - NEVER: </div> <span ...> (space + same line)

   CORRECT — with [INSTRUCTION] in source:
     Q1 "<div class=""qtext"">Welche Geräte sind vorhanden?</div>
<span class=""instruction"">Bitte geben Sie alles Zutreffende an.</span>"
     categorical [0..] {{ _1 "Smartphone" }} ran;

   CORRECT — no [INSTRUCTION] in source:
     Q2 "<div class=""qtext"">Im Folgenden sehen Sie verschiedene Medien. Bitte geben Sie an, wie oft Sie diese privat nutzen.</div>"
     [
         GfKGridType = "banked"
     ]
     loop {{ ... }} ran fields - ( ... ) expand grid;

   CORRECT — instruction embedded at end of question text (Rule C):
     Q17 "<div class=""qtext"">Im Folgenden wollen wir von Ihnen wissen... Mit welchem der folgenden Anbieter verbinden Sie {{#Q17.Loop.Current.Label}} am ehesten?</div>
<span class=""instruction"">Bitte wählen Sie bis zu drei Anbieter aus.</span>"
     [
         GfKGridType = "banked"
     ]
     loop {{ ... }} ran fields - ( ... ) expand grid;

15. DO NOT CHANGE QUESTION / ANSWER MEANING:
   Do not paraphrase unless needed for correct syntax.
   Do not invent providers, genres, sublists, answer categories, texts, or codes.

16. NUMERIC / OE TYPES:
   - If the source clearly says [N], NEVER output numeric;
   - For [N] questions, ALWAYS output text[0..];
   - If the source clearly says [OE], output open text syntax;
   - Keep structure faithful to source.

17. SUBLISTS:
   If the source clearly groups loop items into sublists, preserve them as sublist blocks.
   Example:
   loop
   {{
       sublist1 ""
       {{
           _1 "RTL",
           _2 "VOX"
       }} ran,
       sublist2 ""
       {{
           _3 "SAT.1"
       }} ran
   }} ran fields - ...

18. SIMPLE REPEATED OE LIST QUESTIONS:
   If a question is an open-end question [OE] with multiple blank answer lines like:
   1 __________
   2 __________
   3 __________
   4 __________
   5 __________
   then convert it into a SIMPLE GRID / LOOP WITHOUT BANKED.

   IMPORTANT:
   - Do NOT add [ GfKGridType = "banked" ]
   - Do NOT use categorical
   - Use loop with one statement per blank line
   - Keep the real number of answer lines from the source
   - In fields -, use exactly:
     _scale ""
     text[0..];

19. DO NOT USE MARKDOWN:
   Do NOT wrap output in triple backticks.
   Do NOT add commentary before or after the code.

REFERENCE PATTERNS:

'=====Multiple choice with Base comment + instruction
'Source [BASE] line: "[BASE] Base: all respondents; randomize items except code 99, [M]; [SC];"
'Source question line: "Welche der genannten Geräte...\n[INSTRUCTION] Bitte geben Sie alles Zutreffende an."
'Base: all respondents; randomize items except code 99, [M]; [SC];
Q1 "<div class=""qtext"">Welche der genannten Geräte sind in Ihrem Haushalt vorhanden?</div>
<span class=""instruction"">Bitte geben Sie alles Zutreffende an.</span>"
categorical [0..]
{{
    _1 "Smartphone",
    _99 "keines der genannten" fix exclusive
}} ran;

'=====Grid (banked) WITHOUT instruction - source has no [INSTRUCTION] tag
'Source cell: "Im Folgenden sehen Sie verschiedene Medien. Bitte geben Sie an, wie oft Sie diese privat nutzen."
'No [INSTRUCTION] tag → NO <span class=""instruction""> added
Q2 "<div class=""qtext"">Im Folgenden sehen Sie verschiedene Medien. Bitte geben Sie an, wie oft Sie diese privat nutzen.</div>"
[
    GfKGridType = "banked"
]
loop
{{
    _1 "Klassisches (lineares) Fernsehen",
    _2 "Online-Angebote"
}} ran fields -
(
    _scale ""
    categorical [0..1]
    {{
        _1 "täglich",
        _2 "mehrmals die Woche"
    }};
) expand grid;

'=====Grid (banked) WITH [INSTRUCTION] tag in source
'Source cell: "Bitte bewerten Sie die folgenden Aussagen.\n[INSTRUCTION] Bitte vergeben Sie für jede Zeile eine Bewertung."
Q3 "<div class=""qtext"">Bitte bewerten Sie die folgenden Aussagen.</div>
<span class=""instruction"">Bitte vergeben Sie für jede Zeile eine Bewertung.</span>"
[
    GfKGridType = "banked"
]
loop
{{
    _1 "Aussage A",
    _2 "Aussage B"
}} ran fields -
(
    _scale ""
    categorical [0..1]
    {{
        _1 "stimme voll zu",
        _2 "stimme eher zu"
    }};
) expand grid;

'=====Grid with instruction embedded at end of question text (Rule C)
'Source: "Im Folgenden wollen wir... am ehesten?\n[INSTRUCTION] Bitte wählen Sie bis zu drei Anbieter aus."
Q17 "<div class=""qtext"">Im Folgenden wollen wir von Ihnen wissen, wie sehr Sie die folgenden Sendungsarten / Genres mit bestimmten Streaminganbietern bzw. Mediatheken verbinden. Mit welchem der folgenden Anbieter verbinden Sie {{#Q17.Loop.Current.Label}} am ehesten?</div>
<span class=""instruction"">Bitte wählen Sie bis zu drei Anbieter aus.</span>"
[
    GfKGridType = "banked"
]
loop
{{
    _1 "Serien",
    _2 "Filme"
}} ran fields -
(
    _scale ""
    categorical [0..3]
    {{
        _1 "Netflix",
        _2 "Amazon Prime Video"
    }};
) expand grid;

'=====Example with partial bold and underline only where marked in source (<b> tags in source)
'Source: "<b>Was ist die </b><u><b>maximale</b></u><b> Anzahl...</b>"
Q13 "<div class=""qtext""><strong>Was ist die </strong><u><strong>maximale</strong></u><strong> Anzahl an Anbietern, bei denen Sie sich vorstellen könnten, ein kostenpflichtiges Abo abzuschließen?</strong></div>"
categorical [0..1]
{{
    _1 "1",
    _2 "2"
}};

'=====Example with NO bold in source - plain question text
'Source: "Im Folgenden sehen Sie..." (no <b> tags)
Q15 "<div class=""qtext"">Auf welchen Produktkategorien hat Hofer aktuell zu wenige/unattraktive Aktionen / Angebote?</div>
<span class=""instruction"">Bitte wählen Sie bis zu 5 Produktkategorien aus!</span>"
categorical [0..5]
{{
    _1 "Kategorie 1",
    _2 "Kategorie 2",
    _3 "Kategorie 3"
}} ran;

'=====Repeated OE list example - NO BANKED
Q14 "<div class=""qtext"">Nenne uns deine Lieblingsserien oder -filme oder auch Lieblingssendungen beim Fernsehen und Streaming.</div>
<span class=""instruction"">Du kannst bis zu 5 Titel angeben.</span>"
loop
{{
    _1 "1",
    _2 "2",
    _3 "3",
    _4 "4",
    _5 "5"
}} fields -
(
    _scale ""
    text[0..];
) expand grid;

Raw Survey Document to convert:

{document_text}
""".strip()


# =========================
# FILE UPLOAD
# =========================
uploaded_file = st.file_uploader("Încarcă fișierul .docx", type=["docx"])


# =========================
# MAIN LOGIC
# =========================
if uploaded_file is not None and api_key:
    if st.button("Generează Codul MDD", type="primary"):
        with st.spinner("Citesc documentul și generez scriptul..."):
            try:
                document_text = extract_text_from_docx(uploaded_file)

                if len(document_text) < 10:
                    st.error("Documentul pare să fie gol sau formatul nu a putut fi citit corect.")
                    st.stop()

                genai.configure(api_key=api_key)
                model = genai.GenerativeModel("gemini-2.5-flash")

                prompt = build_prompt(document_text)
                response = model.generate_content(prompt)

                final_code = clean_model_output(get_response_text(response))

                if not final_code:
                    st.error("Modelul nu a returnat text valid.")
                    st.write(response)
                    st.stop()

                st.success("Conversie finalizată!")
                st.code(final_code, language="text")

                st.download_button(
                    label="Descarcă scriptul MDD",
                    data=final_code,
                    file_name="survey_script.mdd",
                    mime="text/plain"
                )

            except Exception as e:
                st.error(f"A apărut o eroare generală: {e}")

elif uploaded_file is not None and not api_key:
    st.warning("Te rog să introduci cheia ")